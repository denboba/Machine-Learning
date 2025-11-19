#!/usr/bin/env python3
"""
Generate Complete Lab Report PDF
Based on tema.pdf requirements and tema1_abdulkadir gobena-denboba.ipynb implementation
"""

import os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors

def create_pdf_report():
    """Create comprehensive PDF report from Word document"""
    
    # Read the Word document
    doc = Document('Assignment1_Complete_Report.docx')
    
    # Create PDF
    pdf_filename = 'Tema_Invatare_Automata_Raport_Complet.pdf'
    pdf = SimpleDocTemplate(
        pdf_filename,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f4788'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1f4788'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#2e5c8a'),
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )
    
    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor('#3d6fa6'),
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
        leading=14
    )
    
    # Process Word document content
    print("Processing Word document content...")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Escape special characters for ReportLab
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Determine style based on paragraph style in Word doc
        style_name = para.style.name
        
        if 'Title' in style_name or text == 'Machine Learning Assignment 1':
            elements.append(Paragraph(text, title_style))
            elements.append(Spacer(1, 0.3*inch))
        elif style_name.startswith('Heading 1') or (
            text.startswith('1.') or text.startswith('2.') or 
            text.startswith('3.') or text.startswith('4.') or 
            text.startswith('5.') or text.startswith('6.')
        ):
            elements.append(Spacer(1, 0.2*inch))
            elements.append(Paragraph(text, heading1_style))
        elif style_name.startswith('Heading 2'):
            elements.append(Paragraph(text, heading2_style))
        elif style_name.startswith('Heading 3'):
            elements.append(Paragraph(text, heading3_style))
        else:
            # Regular text
            try:
                elements.append(Paragraph(text, normal_style))
            except:
                # Skip problematic paragraphs
                pass
    
    # Process tables from Word document
    print(f"Processing {len(doc.tables)} tables...")
    for i, table in enumerate(doc.tables):
        # Convert Word table to ReportLab table
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        
        if data:
            # Create table
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(t)
            elements.append(Spacer(1, 0.2*inch))
    
    # Add available visualizations
    print("Adding visualizations...")
    image_files = [
        ('bike_eda_temporal.png', 'Bike Rental: Temporal Analysis'),
        ('bike_eda_correlations.png', 'Bike Rental: Correlation Analysis'),
        ('bike_eda_distributions.png', 'Bike Rental: Distribution Analysis'),
        ('bike_quantile_regression.png', 'Bike Rental: Quantile Regression'),
        ('bike_models_comparison.png', 'Bike Rental: Model Comparison'),
        ('autovit_eda_target.png', 'Car Prices: Target Variable Analysis'),
        ('autovit_eda_correlations.png', 'Car Prices: Correlation Analysis'),
        ('autovit_eda_categorical.png', 'Car Prices: Categorical Features'),
        ('autovit_eda_year_trend.png', 'Car Prices: Year Trends'),
        ('autovit_eda_missing_pattern.png', 'Car Prices: Missing Data Patterns'),
        ('autovit_models_comparison.png', 'Car Prices: Model Comparison'),
    ]
    
    # Add section for visualizations
    elements.append(PageBreak())
    elements.append(Paragraph('Anexa: Vizualizari', heading1_style))
    elements.append(Spacer(1, 0.2*inch))
    
    for img_file, caption in image_files:
        if os.path.exists(img_file):
            try:
                # Add image with caption
                img = Image(img_file, width=6*inch, height=4*inch, kind='proportional')
                elements.append(img)
                caption_text = caption.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(f'<i>{caption_text}</i>', normal_style))
                elements.append(Spacer(1, 0.3*inch))
            except Exception as e:
                print(f"Could not add image {img_file}: {e}")
    
    # Build PDF
    print(f"Building PDF: {pdf_filename}...")
    pdf.build(elements)
    print(f"✓ PDF report created successfully: {pdf_filename}")
    
    return pdf_filename

if __name__ == '__main__':
    pdf_file = create_pdf_report()
    print(f"\nComplete report generated: {pdf_file}")
    print(f"File size: {os.path.getsize(pdf_file) / 1024:.1f} KB")
